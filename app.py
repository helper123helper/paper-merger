import streamlit as st
from docx import Document
from docx.shared import RGBColor
import io
import re

st.set_page_config(page_title="HMSC 試卷對位工具", page_icon="🎓")

st.title("🎓 HMSC 模擬考：教學檔生成工具 (理想成果版)")
st.info("已修正：自動過濾『提示語(例如:下列任何一項)』，確保答案精準填入表格。")

col1, col2 = st.columns(2)
with col1:
    file_q = st.file_uploader("1. 上傳學生版 (Q)", type=["docx"])
with col2:
    file_ans = st.file_uploader("2. 上傳答案卷 (Ans)", type=["docx"])

if file_q and file_ans:
    if st.button("🪄 生成完美對位教學檔", type="primary"):
        try:
            doc_q = Document(file_q)
            doc_ans = Document(file_ans)
            
            # --- 1. 抓取答案庫 (加入更嚴格的過濾邏輯) ---
            ans_pool = []
            # 定義需要排除的無意義提示語
            exclude_keywords = ["下列任何一項", "建議答案", "題號", "總分", "分數", "分）", "分)"]
            
            for table in doc_ans.tables:
                for row in table.rows:
                    # 抓取該行所有儲存格內容
                    row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
                    for t in row_texts:
                        # 核心過濾：
                        # 1. 長度大於1 (避開 a, b, c)
                        # 2. 不是純數字 (避開 1, 2, 3 分數)
                        # 3. 不包含排除關鍵字
                        if len(t) > 1 and not t.isdigit():
                            if not any(k in t for k in exclude_keywords):
                                # 移除 (x分) 標記
                                clean_text = re.sub(r'[\(（]\s*\d+\s*分\s*[\)）]', '', t).strip()
                                if clean_text:
                                    ans_pool.append(clean_text)
                                    break # 每一列只抓取一個最符合的文字

            ans_idx = 0
            
            # --- 2. 處理題目卷中的表格 (精準填入空白格) ---
            # 這是針對您第一題那種「發展理論表」的填寫
            for table in doc_q.tables:
                for row in table.rows:
                    # 檢查最後一格是否為空，且不是標題列
                    last_cell = row.cells[-1]
                    if ans_idx < len(ans_pool) and not last_cell.text.strip():
                        last_cell.text = ans_pool[ans_idx]
                        # 設置藍色加粗格式
                        for p in last_cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(0, 102, 204)
                        ans_idx += 1

            # --- 3. 處理剩餘的段落題目 (跳過封面) ---
            start_merging = False
            for para in doc_q.paragraphs:
                text = para.text.strip()
                
                # 只有看到「甲部」才開始處理，防止弄壞封面個人資料
                if "甲部" in text:
                    start_merging = True
                if not start_merging:
                    continue
                
                # 偵測題目特徵：結尾有 (x 分)
                if re.search(r'[\(（]\s*\d+\s*分\s*[\)）]$', text):
                    if ans_idx < len(ans_pool):
                        # 避免重複添加
                        if "【建議答案】" not in para.text:
                            run = para.add_run(f"\n【建議答案】：{ans_pool[ans_idx]}")
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0, 102, 204)
                            ans_idx += 1

            # --- 4. 匯出 ---
            output = io.BytesIO()
            doc_q.save(output)
            st.success(f"✅ 對位成功！共填入 {ans_idx} 處內容。")
            st.download_button(
                label="📥 下載理想成果版教學檔",
                data=output.getvalue(),
                file_name="HMSC_教學檔_理想成果.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"系統出錯：{str(e)}")

st.divider()
st.caption("2526 HMSC S6 Mock 專用 - 已解決答案位移問題")
